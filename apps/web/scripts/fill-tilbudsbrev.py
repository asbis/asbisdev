"""Fyll ut kundens tilbudsbrev-mal (Vedlegg 3) med våre opplysninger."""
from docx import Document
from pathlib import Path

SRC = Path("/Users/asbis/code/asbisdev/oppdrag/aktive/2026-105336-rusinfo-selvhjelpsapp/fra-kunden/01-konkurransegrunnlag/03 Vedlegg 3 Mal for tilbudsbrev.docx")
DST = Path("/Users/asbis/code/asbisdev/oppdrag/aktive/2026-105336-rusinfo-selvhjelpsapp/til-innsending/autogenerert/1.1 Tilbudsbrev Asbjørn Rørvik.docx")

FIELDS = {
    "Firmanavn: xxxx": "Firmanavn: Asbjørn Rørvik (enkeltpersonforetak)",
    "Org.nr: xxxx": "Org.nr: 820 252 632",
    "Adresse: xxxx": "Adresse: Vipeveien 7B",
    "Poststed: xxxx": "Poststed: 4323 Sandnes",
    "Dato: dd.mm.åååå": "Dato: 24.04.2026",
}

AVVIK_TEXT = (
    "Ett avvik registrert: Leverandøren tilbyr alternativ løsning på krav 4.2.1 og 4.2.3 "
    "i Bilag 1 — React Native + TypeScript i stedet for Ionic 7 + Vue 3 som ønsket rammeverk. "
    "Appens struktur, informasjonsarkitektur, skjermhierarki, datamodell og backend-modell (Strapi) "
    "er speilet 1:1 med HAP — kun renderingslaget endres (fra webview til native UI). "
    "Begrunnelse: bedre brukeropplevelse for en sårbar målgruppe, som må være det endelige målet, "
    "kombinert med bredere vedlikeholdsmarked i Norge og modnere tilgjengelighets-API. "
    "Fullstendig redegjørelse og konsekvensvurdering fremgår av Bilag 2 § 4.2.3 "
    "(i utfylt Bilag 2 Leverandørens løsningsspesifikasjon)."
)

JA_NEI_TEXT = "JA [X]   NEI [ ]"

# Kontaktinfo-tabell
CONTACT = {
    "Navn:": "Asbjørn Rørvik",
    "Telefonnummer:": "+47 47 65 86 51",
    "Epost:": "hei@asbjornrorvik.dev",
}

NAME_SIGNATURE = "Navn: ASBJØRN RØRVIK (i blokkbokstaver)"
STILLING_SIGNATURE = "Stilling: Innehaver / daglig leder"


def replace_paragraph_text(p, new_text: str, *, bold: bool = False):
    """Erstatt tekst i en paragraf uten å miste formatering."""
    # Tøm alle run
    for r in p.runs:
        r.text = ""
    if p.runs:
        p.runs[0].text = new_text
        if bold:
            p.runs[0].bold = True
    else:
        run = p.add_run(new_text)
        if bold:
            run.bold = True


def main():
    doc = Document(SRC)

    for p in doc.paragraphs:
        t = p.text.strip()
        # Header-felter
        if t in FIELDS:
            replace_paragraph_text(p, FIELDS[t])
            print(f"  ✓ {t[:30]} → utfylt")
            continue

        # Avvik-felt
        if t.startswith("Eventuelle avvik"):
            replace_paragraph_text(p, "Eventuelle avvik: " + AVVIK_TEXT)
            print(f"  ✓ Avvik beskrevet")
            continue

        # Ja/Nei-kryss
        if "(sett kryss)" in t:
            replace_paragraph_text(p, JA_NEI_TEXT + "  (vedlegg medfølger, jf. konkurransegrunnlagets punkt 7)")
            print(f"  ✓ JA-kryss satt")
            continue

        # Navn i blokkbokstaver
        if t.startswith("Navn: xxxx"):
            replace_paragraph_text(p, NAME_SIGNATURE)
            print(f"  ✓ Signaturnavn utfylt")
            continue

        # Stilling
        if t.startswith("Stilling: xxxx"):
            replace_paragraph_text(p, STILLING_SIGNATURE)
            print(f"  ✓ Stilling utfylt")
            continue

    # Fyll ut kontakt-tabellen
    for tbl in doc.tables:
        for row in tbl.rows:
            if len(row.cells) >= 2:
                label = row.cells[0].text.strip()
                if label in CONTACT:
                    # Erstatt xxxx
                    cell = row.cells[1]
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.text = ""
                    if cell.paragraphs and cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].text = CONTACT[label]
                    else:
                        cell.paragraphs[0].add_run(CONTACT[label])
                    print(f"  ✓ {label} → {CONTACT[label]}")

    DST.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DST)
    print(f"\nLagret til: {DST}")


if __name__ == "__main__":
    main()

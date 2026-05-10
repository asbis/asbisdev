"""Erstatt placeholdere i kundens Bilag 3-10 med referanse til vår prosjektplan."""
from docx import Document
from pathlib import Path

SRC = Path("/Users/asbis/code/asbisdev/oppdrag/aktive/2026-105336-rusinfo-selvhjelpsapp/fra-kunden/02-ssa-t-kontrakt/03 Del 2 SSA-T Bilag 3 til 10.docx")
DST = Path("/Users/asbis/code/asbisdev/oppdrag/aktive/2026-105336-rusinfo-selvhjelpsapp/til-innsending/autogenerert/03 Del 2 SSA-T Bilag 3 til 10 (utfylt).docx")

REPLACEMENT = "Besvares i eget vedlegg: «04 Bilag 4 — Prosjektplan.pdf» (vedlagt tilbudet)."

doc = Document(SRC)
count = 0
for p in doc.paragraphs:
    if "<Besvares i eget vedlegg til bilag 4>" in p.text:
        # Erstatt hele paragrafens tekst
        for run in p.runs:
            run.text = ""
        if p.runs:
            p.runs[0].text = REPLACEMENT
            p.runs[0].italic = True
        else:
            p.add_run(REPLACEMENT).italic = True
        count += 1
        print(f"  ✓ Erstattet placeholder ({count})")

DST.parent.mkdir(parents=True, exist_ok=True)
doc.save(DST)
print(f"\n{count} placeholdere erstattet. Lagret til: {DST}")

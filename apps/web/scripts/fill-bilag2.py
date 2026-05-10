"""Fyll ut kundens Word-mal for Bilag 2 med våre besvarelser."""
from docx import Document
from docx.shared import Pt
from pathlib import Path
import sys

SRC = Path("/Users/asbis/code/asbisdev/oppdrag/aktive/2026-105336-rusinfo-selvhjelpsapp/fra-kunden/02-ssa-t-kontrakt/02 DEL 2 SSA-T Bilag 2 Leverandørens løsningsspesifikasjon.docx")
DST = Path("/Users/asbis/code/asbisdev/oppdrag/aktive/2026-105336-rusinfo-selvhjelpsapp/til-innsending/autogenerert/02 DEL 2 SSA-T Bilag 2 — Leverandørens løsningsspesifikasjon (utfylt).docx")

# Besvarelser per krav. Nøkkel matcher eksakt krav-nr fra malen.
BESVARELSER = {
    "4.1.1": """Lest og akseptert.

Appen leveres for iOS (15.0+) og Android (8.0 / API 26+), utviklet i React Native med delt TypeScript-kildekode. Native moduler benyttes der hardware-tilgang eller plattform-spesifikk UX krever det (biometri, haptikk, lokal oversettelse).

Krav til Kundens tekniske plattform:
• Én webserver (Linux, 2 vCPU / 4 GB RAM) for Strapi CMS + Postgres
• HTTPS med gyldig sertifikat (Let's Encrypt eller Oslo kommunes interne CA)
• Apple Developer Program-konto og Google Play Console-konto eid av RUSinfo/Velferdsetaten
• Git-hosting (GitHub, GitLab eller Azure DevOps) for kildekode-eierskap

Leverandør bistår med oppsett.""",

    "4.1.2": """Ønsket løsning kan leveres uten avvik.

Vedlikeholdsavtalen følger SSA-T og dekker 3 år med automatisk 1-årig fornyelse, oppsigbar med 3 måneders varsel.

Innhold:
• Feilretting og feilsøking i app + backend-server
• Sikkerhetsoppdateringer (React Native-versjoner, Expo-oppgraderinger, Strapi-patches, alle tredjepartsbiblioteker)
• Årlige plattformoppgraderinger (iOS SDK, Android target-API) for å beholde butikk-godkjenning
• Drift og overvåking av Strapi-server (helsesjekk, backup, logg-rotasjon)
• Innmelding av feil direkte fra appen via innebygget "send feedback"-kanal som posterer anonymt til Strapi
• Push av innholdsendringer via Strapi uten nytt app-release
• Responstid: 24 timer på virkedager — bedre enn ønsket 48 timer

Videreutvikling av nye funksjoner prises separat etter Bilag 7. Oppdragsgiver får skriftlig estimat (timer × timepris) og må godkjenne før arbeid starter.""",

    "4.1.3": """Lest og akseptert.

Appen bruker ingen generativ KI i sluttproduktet. Lokal oversettelsesmodell (jf. krav 4.2.3) er ferdig trent av plattformleverandør (Apple/Google) før modellen pakkes med appen. Modellen trenes ikke på Kundens eller brukernes data.

KI-assistent (Claude/Copilot) brukes i utviklingsfasen som kodeverktøy — aldri som del av produktets kjøretidslogikk. Se 4.1.4.""",

    "4.1.4": """Ønsket løsning kan leveres uten avvik.""",

    "4.1.5": """Lest og akseptert.

Opplæring tilbys i tre former:
• Oppstarts-workshop (4 timer, fysisk eller Teams): gjennomgang av Strapi-admin — innholdsredigering, dagens tema, default kriseplan, bildebibliotek, statistikkhenting
• Skriftlig admin-manual (PDF + Markdown i git, med skjermbilder) levert ved go-live
• Teknisk bistand i vedlikeholdsperioden via e-post/chat, timebasert etter Bilag 7

Priset separat i Bilag 7.""",

    "4.1.6": """Lest og akseptert.

Leverandør håndterer hele publiseringsløpet i Oppdragsgivers egne Apple Developer- og Google Play-kontoer:
• App Store Connect: metadata, screenshots, beskrivelse, aldersvurdering, eksport-compliance
• Google Play Console: metadata, screenshots, Data Safety-skjema, policy-samsvar
• Håndtering av review-prosessen og svar på eventuelle avslag
• Versjonsnummerering og release-noter

Leverandør har tidligere vært med på publisering til App Store og Google Play gjennom arbeid hos Netpower (Kolumbus, NLR TryggDrift) og Bouvet (Easee, Equinor), og har i tillegg egenhendig publisert og drifter Supportify-integrasjoner.""",

    "4.1.7": """Lest og akseptert.

To manualer leveres ved go-live:
• Sluttbrukermanual (PDF + web-publisering på rusinfo.no): illustrert innføring modul for modul, inkludert hvordan bruker aktiverer lås, hvordan kriseplanen fungerer, hvordan motivasjonsveggen bygges
• Admin-manual for RUSinfo-personell (PDF + Markdown i git): Strapi-redigering, statistikkhenting, feilhåndtering, release-rutiner

Begge versjoneres i git slik at nye versjoner alltid ligger oppe ved app-oppdateringer.""",

    "4.1.8": """Lest og akseptert.

Hele appens kildekode er Oppdragsgivers eiendom. Praksis:
• Git-repo opprettet hos Oppdragsgiver (GitHub, GitLab eller Azure DevOps) fra dag 1. Leverandør committer direkte til Kundens repo — ingen mellomlagring.
• Ved kontraktsslutt har Oppdragsgiver fullstendig tilgang, eierskap og rett til videreutvikling, eventuelt med ny leverandør.
• Lisens: Oppdragsgiver velger (MIT/Apache-2.0 for åpen kildekode, eller lukket proprietær).
• Alle tredjepartsavhengigheter listes i Bilag 10 med lisens og versjon. Kun lisenser som er kompatible med offentlig bruk benyttes (MIT, BSD, Apache-2.0, ISC).""",

    "4.1.9": """Lest og akseptert.

Appen inneholder:
• Ingen reklame (ingen AdMob, ingen banner, ingen sponset innhold)
• Ingen tracking-pixler eller tredjeparts analytikk (Google Analytics, Firebase Analytics, Segment, Mixpanel — ingen av disse)
• Ingen remarketing eller attribution-SDK-er

Eneste telemetri er frivillig, anonym statistikk til RUSinfos egen Strapi (ingen eksterne mottakere).""",

    "4.1.10": """Lest og akseptert.

Appen leveres på norsk bokmål. All brukerrettet tekst er sentralisert i src/i18n/no.ts, som gjør lokalisering til nynorsk/engelsk/andre språk trivielt senere dersom RUSinfo ønsker det.

Innhold som RUSinfo styrer fra Strapi (artikler, ukens tema, default kriseplan) administreres på norsk i admin-UI; flerspråklig Strapi-struktur er forberedt i datamodellen for framtidig utvidelse.""",

    "4.1.11": """Lest og akseptert. Oversikt over reduserbare komponenter med estimert priseffekt (NOK ekskl. mva):

• Kartleggingsverktøy (kalender, planlagt vs. faktisk): ~60 000 — fjerner planleggingsmodulen, tidsbruk-logging beholdes i dagboken
• Motivasjonsvegg (galleri + komposisjon): ~35 000 — personliggjøring via galleri fjernes, standard-bakgrunner beholdes
• Bildebibliotek (kuratert bakgrunnssett): ~20 000 — standard bakgrunner i stedet for valgbare
• Lokal AI-oversettelse: ~30 000 — app kun på norsk uten on-device oversetting
• Palett-velger utover grayscale: ~20 000 — ett tema + grayscale. Grayscale beholdes alltid (UU-krav).
• Krise-chat (Kord AI anonym samtale): ~40 000 — kriseplan beholdes, chat-funksjonen utgår
• Oppstarts-workshop: ~15 000 — kun skriftlig manual + videoopptak

Disse er uavhengige og kan tas ut i kombinasjon. Kjernefunksjoner (onboarding, sikkerhet, dagbok, oversikt, prestasjoner, info, kriseknapp, default kriseplan, publisering, manualer) kan ikke tas ut uten å bryte oppdragets overordnede behov.""",

    "4.2.1": """Leverandøren tilbyr en alternativ løsning som oppfyller Kundens overordnede behov, men som avviker delvis fra ønsket løsning.

⏵ Prøv fungerende prototype i nettleseren: https://asbjornrorvik.dev/no/oppdrag/rusinfo-app (passord: rusinfo2026)

Avvik: rammeverk (React Native i stedet for Ionic 7/Vue 3). Struktur, IU, XU, skjermhierarki, datamodell og backend-modell er speilet 1:1 med HAP — kun renderingslaget endres fra webview til native UI.

Begrunnelse: se krav 4.2.3 (rammeverkvalg).

Konsekvens for IU/XU/backend: ingen. Brukeropplevelsen er mer native, ikke annerledes.

Faktisk oppbygging (implementert i vedlagt prototype):
• Onboarding: flertrinns flyt — modul → startdato → valgfri demografisk avklaring (aldersgruppe, kjønn, fylke, rolle [bruker/pårørende/profesjonell]) som kun brukes til anonym statistikk → frekvens + kostnad → personvernløfte → sikkerhetsvalg → PIN → nødkontakt
• Hovednavigasjon: 5 tabs — Hjem, Dagbok, Oversikt, Prestasjon, Info
• Modaler: native-stack-skjermer for Motivasjonsvegg, Kartlegging, Artikkel, Dagboksinnlegg, Innstillinger, Chat, Kriseplan som bottom sheet
• Global kriseknapp: flytende FAB med pust-animasjon, synlig på alle hovedskjermer
• Backend: Strapi headless CMS, selv-hostet hos RUSinfo, innholdstyper speiler HAP

Anonym statistikk: demografisk avklaring i onboarding (alder, kjønn, fylke, rolle) lagres aggregert og ikke-koblbart til bruker, kun for å gi RUSinfo demografisk oversikt over brukerne av appen.

Universell utforming (WCAG 2.1 AA + EUs webdirektiv):
• accessibilityLabel, accessibilityRole, accessibilityState på alle interaktive komponenter (VoiceOver/TalkBack)
• Kontrast ≥ 4.5:1 for normal tekst, ≥ 3:1 for stor tekst
• Grayscale-modus som alternativt tema
• Dynamic Type-støtte og Reduce Motion-respekt
• Full tastatur- og skjermleser-navigasjon testet

Ekstern UU-revisjon utføres av Funka Nu eller MediaLT før publisering.""",

    "4.2.2": """Ønsket løsning kan leveres uten avvik.

⏵ Prøv fungerende prototype i nettleseren: https://asbjornrorvik.dev/no/oppdrag/rusinfo-app (passord: rusinfo2026)""",

    "4.2.3": """Leverandøren tilbyr en alternativ løsning som oppfyller Kundens overordnede behov, men som avviker fra ønsket løsning.

Felles kodebase: Ja. Én React Native + TypeScript-kodebase kompileres til både iOS og Android uten forretningslogikk-duplisering. Plattform-spesifikk kode er isolert til native modules (biometri, oversettelse, haptikk).

Rammeverklevetid 5–7 år: Oppfylt med god margin. React Native er utviklet av Meta siden 2015, aktivt utviklet (månedlig releases), og brukes i produksjon av Instagram, Shopify, Discord, Microsoft Office, Coinbase — samt mange norske produksjonsapper (DNB, Vy, Equinor, Easee). Expo har garantert support-syklus.

Backend-server: Ja, samme modell som HAP — headless CMS (Strapi) som kommuniserer med app over HTTPS-JSON-API. RUSinfo får samme type admin-UI for å redigere innhold og hente ut statistikk. Selv-hostes i Oppdragsgivers driftsmiljø.

Lydfiler: Ja. React Native støtter lydavspilling via expo-av / react-native-sound. Lyd-assets kan være lokale i app-bundle eller strømmes fra Strapi.

Lokal oversettelsesmodell på enheten: oversettelse gjelder kun faktatekster produsert av RUSinfo i backend-løsningen (artikler, info-innhold, ukens tema, default kriseplan) — ikke brukerens egeninndata (dagboksinnlegg, motivasjonsvegg-tekst). Innholdet oversettes fra norsk til utvalgte språk slik at informasjonen gjøres tilgjengelig også for brukere uten norsk som morsmål.

• iOS 17.4+: Apple Translation Framework — helt on-device, gratis, Apple-vedlikeholdt, dekker norsk
• Android: Google ML Kit Translation API — on-device, gratis, Google-vedlikeholdt
• Fallback for eldre iOS: pakket Helsinki-NLP Marian-modell via ONNX Runtime (~30 MB)

Fordelen: null vedlikeholdsbyrde, null datatrafikk, null risiko for leverandøruavhengighet.

Statistikkinnsamling: Som HAP — frivillig, anonym innsending til Strapi (aggregert, ikke-koblbar til person).

Avvik fra "samme rammeverk som HAP":
• HAP bruker Ionic 7 + Vue 3 (Capacitor-basert webvisning i native container)
• Vi bruker React Native + Expo (faktisk native UI-komponenter)

Begrunnelse:
1. Brukeropplevelse for sårbar målgruppe. React Native rendrer til iOS UIKit / Android Material — native animasjoner og gesture-håndtering. For en bruker i krise gir dette smidigere interaksjon, særlig på eldre enheter.
2. Bredere vedlikeholdsmarked i Norge. React Native-utviklere er vesentlig flere enn Ionic/Vue-utviklere, noe som reduserer Oppdragsgivers risiko for leverandør-lock-in.
3. Modnere tilgjengelighets-API. React Natives accessibility-prop er direkte koblet til plattformens egne skjermleser-API-er, uten webview-mellomlag.
4. Større biblioteksøkosystem via Expo-teamet.

HAP-paritet bevares der det teller: skjermhierarki, informasjonsarkitektur, datamodell, innholdstyper i Strapi og admin-opplevelsen for RUSinfo er speilet 1:1 med HAP. Redaktører som kjenner HAPs admin kan bruke den nye appens backend uten omlæring, og innhold kan migreres direkte mellom systemene. Det som endres er utelukkende renderingslaget.

Konsekvens: kildekoden bygges fra bunnen i React Native (kan ikke porteres direkte fra Ionic/Vue). Dette er inkludert i hovedoppdraget uten ekstra kostnad. For RUSinfo betyr valget at sluttbrukeren — som er det endelige målet — får en app som føles native, responsiv og rolig, mens dere beholder friheten til senere å gjenbruke innhold, struktur og backend-modell fra HAP.

Ingen konsekvens for funksjonalitet, personvern, brukerflyt, innhold, publiseringsprosess eller driftskrav.""",

    "4.3.1": """Ønsket løsning kan leveres uten avvik.""",

    "4.4.1": """Lest og akseptert.

Siden løsningen (jf. § 4.3.1) ikke behandler personopplysninger, er krav §§ 4.4.1–4.4.13 i utgangspunktet ikke utløst. Leverandør bekrefter likevel oppfyllelse for det tilfelle at Kunden senere velger å introdusere slik behandling.

Standard DFØ-databehandleravtale inngås samtidig med den endringen som utløser behovet.""",

    "4.4.2": """Lest og akseptert.

Leverandør har erfaring med personvern- og compliance-arbeid gjennom tidligere teamarbeid på Easee EaseePay (PCI DSS-miljø, via Bouvet 2021–2022) og har etablerte interne rutiner for sikker kode, tilgangsstyring og dataminimering. Dokumentasjon leveres på forespørsel.""",

    "4.4.3": """Lest og akseptert.

Leverandør bistår med risikovurdering (DPIA-forberedelse, trusselmodellering) ved etablering og drift dersom fremtidige endringer utløser behov for behandling av personopplysninger.""",

    "4.4.4": """Lest og akseptert.

Eneste planlagte underleverandør er Funka Nu eller MediaLT for ekstern UU-revisjon. Underleverandør godkjennes skriftlig av Kunden før bruk. Underleverandør får ikke tilgang til personopplysninger.""",

    "4.4.5": """Ikke relevant — løsningen behandler ikke personopplysninger jf. § 4.3.1. Leverandør følger prinsippene i ISO/IEC 27001 for sikkerhetsstyring i egen virksomhet. Bekreftes oppfylt dersom framtidige endringer utløser krav om formelt styringssystem.""",

    "4.4.6": """Ikke relevant jf. § 4.3.1. ROS-analyse utføres dersom framtidige endringer introduserer behandling av personopplysninger, i samarbeid med Kunden.""",

    "4.4.7": """Ikke relevant jf. § 4.3.1 — løsningen har ingen brukerkontoer eller persondata. Leverandørens git-repo-tilgang og Strapi-admin følger prinsippet om minste nødvendige tilgang og revokeres ved avtaleslutt.""",

    "4.4.8": """Strapi-admin benytter sterke passord og valgfri MFA. Sikker lagring av hemmeligheter via Azure Key Vault eller tilsvarende, eid og administrert av Kunden. Tilgang følger minste nødvendige-prinsippet. Ikke relevant for persondata jf. § 4.3.1.""",

    "4.4.9": """All kommunikasjon mellom app og Strapi er kryptert via HTTPS (TLS 1.2+). Strapi-server kjører bak reverse proxy med gyldig sertifikat. Ingen persondata transporteres jf. § 4.3.1.""",

    "4.4.10": """Ikke relevant jf. § 4.3.1 — det finnes ingen persondata å sikre. Innholdsdata i Strapi (artikler, temaer) er uten personopplysninger og sikres via standard Strapi-mekanismer + daglig database-backup.""",

    "4.4.11": """Sertifikater (TLS for Strapi, App Store Connect, Google Play-signering) håndteres av Kunden i eget hemmelighetslager. Leverandør får tilgang under prosjektet og revokeres automatisk ved avtaleslutt.""",

    "4.4.12": """Strapi-database sikkerhetskopieres daglig til Kundens driftsmiljø. Retensjonstid avtales med Kunden (anbefalt: 30 dager). App-kildekoden ligger i git hos Kunden og er i seg selv en kontinuerlig backup.""",

    "4.4.13": """Ved sikkerhetshendelse varsles Kunden skriftlig innen 24 timer. Leverandør samarbeider med Kundens personvernombud og eventuelt Datatilsynet dersom varslingsplikt utløses. Rutiner etableres ved kontraktsstart.""",
}


def fill_besvarelse_cell(cell, text: str):
    """Erstatt celleinnhold, beholder "Leverandørens besvarelse:" som header."""
    # Rens eksisterende tekst
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""

    # Første paragraf: header (fet)
    p0 = cell.paragraphs[0]
    run = p0.add_run("Leverandørens besvarelse:")
    run.bold = True

    # Legg til besvarelse som nye paragrafer
    for line in text.split("\n"):
        p = cell.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(10)


def main():
    if not SRC.exists():
        print(f"Fant ikke: {SRC}", file=sys.stderr)
        sys.exit(1)

    doc = Document(SRC)
    filled = 0
    missing = []

    for tbl in doc.tables:
        rows = tbl.rows
        for r_idx, row in enumerate(rows):
            # Besvarelses-rad: alle 3 cellene starter med "Leverandørens besvarelse:"
            cell_texts = [c.text.strip() for c in row.cells]
            if all(t.startswith("Leverandørens besvarelse:") for t in cell_texts):
                if r_idx == 0:
                    continue
                prev = rows[r_idx - 1].cells[0].text.strip()
                krav_nr = None
                for line in prev.split("\n"):
                    line = line.strip()
                    # Match 4.1.1, 4.2.3, 4.4.13 osv.
                    parts = line.split(".")
                    if len(parts) >= 3 and all(p.isdigit() for p in parts[:3]):
                        krav_nr = f"{parts[0]}.{parts[1]}.{parts[2]}"
                        break
                    if len(parts) == 2 and all(p.isdigit() for p in parts):
                        krav_nr = f"{parts[0]}.{parts[1]}"
                        break
                if not krav_nr:
                    continue
                if krav_nr in BESVARELSER:
                    fill_besvarelse_cell(row.cells[0], BESVARELSER[krav_nr])
                    filled += 1
                    print(f"  ✓ {krav_nr}")
                else:
                    missing.append(krav_nr)

    DST.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DST)
    print(f"\n{filled} besvarelser fylt inn")
    if missing:
        print(f"Manglet besvarelse for: {missing}")
    print(f"Lagret til: {DST}")


if __name__ == "__main__":
    main()

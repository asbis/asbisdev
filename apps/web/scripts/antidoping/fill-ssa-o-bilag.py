"""Fyll ut SSA-O Bilag (samle-fil) — Bilag 1, 2, 3, 4, 5, 7 m.fl.

Kundens fil `2.1 SSA-O Bilag.docx` inneholder ALLE bilag i én fil. Vi fyller:
  - Bilag 2: Konsulentens spesifikasjon — 4 kvalitets­kriterier (Tabell 0)
  - Bilag 1.5 / nøkkelressurser-tabeller (Tabell 1, 3)
  - Bilag 7: Pristabeller (Tabell 4, 5, 6, 7)

Besvarelsene for de 4 kvalitets­kriteriene leses fra de fire markdown-filene i
`vart-utkast/`. Det gjør at innholdet kun finnes ett sted.
"""
from docx import Document
import re
import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent))
from _paths import SRC_DIR, DST_DIR, UTKAST_DIR
from _docx_utils import set_cell_text, render_markdown_to_cell

SRC = SRC_DIR / "2.1 SSA-O Bilag.docx"
DST = DST_DIR / "SSA-O Bilag (utfylt).docx"


def load_md_body(path: Path) -> str:
    """Les markdown, fjern frontmatter, returner kropp."""
    text = path.read_text(encoding="utf-8")
    # Strip YAML frontmatter (--- ... ---)
    text = re.sub(r"^---\n.*?\n---\n", "", text, count=1, flags=re.DOTALL)
    return text.strip()


# Hver av de fire kvalitets­kriteriene fra Bilag 2 (konkurransegrunnlag pkt. 7.1):
KRITERIER = {
    1: UTKAST_DIR / "bilag2-1-app-losning.md",          # App-løsningen og dens kvalitet
    2: UTKAST_DIR / "bilag2-2-nokkelressurser.md",       # Nøkkelressurser
    3: UTKAST_DIR / "bilag2-3-oppdragsforstaelse.md",    # Oppdragsforståelse
    4: UTKAST_DIR / "bilag2-4-service-support.md",       # Service, support og vedlikehold
}

# Pris (Bilag 7) — VELG ÉN modell og oppdater her før kjøring.
PRIS_FASTPRIS_EKS = 1_280_000   # NOK ekskl. mva (=1 600 000 inkl. 25 % mva)
PRIS_FASTPRIS_INKL = 1_600_000  # NOK inkl. mva — matcher BILAG-2-SAMLET.md
PRIS_TIME_EKS = 1_450           # NOK/t ekskl. mva
PRIS_TIME_INKL = 1_812          # NOK/t inkl. 25 % mva (1 450 × 1.25 = 1 812.50, rundes)
TOTALPRIS_EKS = PRIS_FASTPRIS_EKS  # Justeres hvis time-anslag legges på toppen
TOTALPRIS_INKL = PRIS_FASTPRIS_INKL

NOKKELRESSURSER = [
    {"navn": "Asbjørn Rørvik", "kategori": "Senior fullstack-utvikler / prosjektleder", "kompetanse": "React Native, Flutter, TypeScript, Go, native iOS/Android publisering"},
]

# Underleverandører i kvalifikasjonsforstand: ingen.
# UU-revisjon kjøpes som ekstern tjeneste (Funka Nu / MediaLT) og er priset inn i fastpris.
EGNE_UNDERLEVERANDORER = []


def fill_kvalitetskriterier(doc):
    """Tabell 0: besvarelse-radene under hvert kriterium 1–4.

    Mal-strukturen er ikke helt konsistent — derfor bruker vi kjente rad-indekser
    for hvert kriterium (verifisert mot 2.1 SSA-O Bilag.docx, 15 rader):
      R4  = besvarelse kriterium 1
      R8  = besvarelse kriterium 2
      R11 = besvarelse kriterium 3
      R14 = besvarelse kriterium 4
    """
    t = doc.tables[0]
    rows = t.rows
    filled = []

    rad_per_kriterium = {1: 4, 2: 8, 3: 11, 4: 14}

    for nr, ridx in rad_per_kriterium.items():
        if ridx >= len(rows):
            print(f"  ⚠ Tabell 0 har bare {len(rows)} rader — kriterium {nr} mangler")
            continue
        md_path = KRITERIER[nr]
        if not md_path.exists():
            print(f"  ⚠ Mangler markdown: {md_path}")
            continue
        body = load_md_body(md_path)
        render_markdown_to_cell(rows[ridx].cells[1], body, size=10)
        filled.append(nr)
    return filled


def fill_nokkelressurser(doc):
    """Tabell 1 (4x2 Navn/Kategori) — egne nøkkelressurser."""
    if len(doc.tables) <= 1:
        return
    t = doc.tables[1]
    # Rad 0 = header, rad 1+ = data-rader
    for idx, person in enumerate(NOKKELRESSURSER):
        if idx + 1 >= len(t.rows):
            break
        set_cell_text(t.cell(idx + 1, 0), person["navn"])
        set_cell_text(t.cell(idx + 1, 1), person["kategori"])


def fill_underleverandorer(doc):
    """Tabell 3 (6x3 Navn/Kategori/Kompetanse) — egne underleverandører."""
    if len(doc.tables) <= 3:
        return
    t = doc.tables[3]
    for idx, ul in enumerate(EGNE_UNDERLEVERANDORER):
        if idx + 1 >= len(t.rows):
            break
        set_cell_text(t.cell(idx + 1, 0), ul["navn"])
        set_cell_text(t.cell(idx + 1, 1), ul["kategori"])
        if "kompetanse" in ul:
            set_cell_text(t.cell(idx + 1, 2), ul.get("kompetanse", ""))


def fill_priser(doc):
    """Tabell 4 (Pris for Oppdraget), 5 (Pris per time), 6 (Totalpris)."""
    if len(doc.tables) <= 6:
        return
    fastpris = doc.tables[4]
    # R1: Pris for Oppdraget ekskl. mva
    set_cell_text(fastpris.cell(1, 1), "NOK")
    set_cell_text(fastpris.cell(1, 2), f"{PRIS_FASTPRIS_EKS:,}".replace(",", " "))
    set_cell_text(fastpris.cell(2, 1), "NOK")
    set_cell_text(fastpris.cell(2, 2), f"{PRIS_FASTPRIS_INKL - PRIS_FASTPRIS_EKS:,}".replace(",", " "))
    # Mva-prosent
    fastpris.cell(2, 0).text = "Mva. 25 %"
    # R3: Kontraktssum (allerede satt til 1 600 000 i mal — vi overskriver med vår verdi)
    set_cell_text(fastpris.cell(3, 2), f"{PRIS_FASTPRIS_INKL:,}".replace(",", " "))

    timepris = doc.tables[5]
    set_cell_text(timepris.cell(1, 1), "NOK")
    set_cell_text(timepris.cell(1, 2), f"{PRIS_TIME_EKS:,}".replace(",", " "))
    set_cell_text(timepris.cell(2, 1), "NOK")
    set_cell_text(timepris.cell(2, 2), f"{PRIS_TIME_INKL - PRIS_TIME_EKS:,}".replace(",", " "))
    timepris.cell(2, 0).text = "Mva. 25 %"
    set_cell_text(timepris.cell(3, 1), "NOK")
    set_cell_text(timepris.cell(3, 2), f"{PRIS_TIME_INKL:,}".replace(",", " "))

    totalpris = doc.tables[6]
    set_cell_text(totalpris.cell(1, 1), "NOK")
    set_cell_text(totalpris.cell(1, 2), f"{TOTALPRIS_EKS:,}".replace(",", " "))
    set_cell_text(totalpris.cell(2, 1), "NOK")
    set_cell_text(totalpris.cell(2, 2), f"{TOTALPRIS_INKL:,}".replace(",", " "))


def main():
    doc = Document(SRC)

    filled_kriterier = fill_kvalitetskriterier(doc)
    fill_nokkelressurser(doc)
    fill_underleverandorer(doc)
    fill_priser(doc)

    DST.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DST)
    print(f"✓ Lagret: {DST}")
    print(f"  Kvalitets­kriterier fylt: {filled_kriterier} (forventet [1, 2, 3, 4])")
    print("  ⚠ Manuelle steg:")
    print("    - Verifiser pristabellene (Bilag 7) — diskrepans mellom analyse (1,45M+pool) og samlet (1,6M)")
    print("    - Fyll inn betalingsplan-tabellen hvis ADNO ber om justering")
    print("    - Tabell 3 (egne underleverandører): tom — UU-revisjon er ekstern tjeneste, ikke underleverandør")


if __name__ == "__main__":
    main()

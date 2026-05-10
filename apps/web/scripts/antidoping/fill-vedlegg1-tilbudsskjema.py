"""Fyll ut Vedlegg 1 — Tilbudsskjema for antidoping-app."""
from docx import Document
import sys
sys.path.insert(0, str(__import__("pathlib").Path(__file__).parent))
from _paths import SRC_DIR, DST_DIR, LEVERANDOR
from _docx_utils import set_cell_text, fill_label_value

SRC = SRC_DIR / "1.1 Vedlegg 1 - Tilbudsskjema.docx"
DST = DST_DIR / "Vedlegg 1 — Tilbudsskjema (utfylt).docx"


def main():
    doc = Document(SRC)

    # Tabell 0: Opplysninger om Leverandøren
    t0 = doc.tables[0]
    fill_label_value(t0, "Navnet på virksomheten", LEVERANDOR["navn"])
    fill_label_value(t0, "Organisasjonsnummer", LEVERANDOR["orgnr"])
    fill_label_value(t0, "Postadresse", LEVERANDOR["postadresse"])
    fill_label_value(t0, "Postnummer", LEVERANDOR["postnummer"])
    fill_label_value(t0, "Sted", LEVERANDOR["sted"])

    # Tabell 1: Kontaktperson
    t1 = doc.tables[1]
    fill_label_value(t1, "Navn", LEVERANDOR["kontakt_navn"])
    fill_label_value(t1, "Stilling", LEVERANDOR["kontakt_stilling"])
    fill_label_value(t1, "E-postadresse", LEVERANDOR["kontakt_epost"])
    fill_label_value(t1, "Telefonnummer", LEVERANDOR["kontakt_telefon"])

    # Tabell 2: Signatar
    t2 = doc.tables[2]
    fill_label_value(t2, "Navn", LEVERANDOR["signatar_navn"])
    fill_label_value(t2, "Stilling", LEVERANDOR["signatar_stilling"])
    # Fullmakt-rad: avkrysning gjøres manuelt før signering

    # Tabell 6 (Erklæring): sted/dato
    t6 = doc.tables[6]
    fill_label_value(t6, "Sted og dato", f"Stavanger, [DATO — fylles ut ved signering]")

    DST.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DST)
    print(f"✓ Lagret: {DST}")
    print("  ⚠ Manuelle steg før innsending:")
    print("    - Kryss av for taushetsplikt (Tabell 3)")
    print("    - Kryss av for avvik=Nei (Tabell 4)")
    print("    - Kryss av for leverandørkonstellasjon (Tabell 5) — 'Én leverandør alene' (UU-revisjon kjøpes som ekstern tjeneste, ikke som underleverandør i kvalifikasjonsforstand)")
    print("    - Fyll dato + signer elektronisk i Mercell")


if __name__ == "__main__":
    main()

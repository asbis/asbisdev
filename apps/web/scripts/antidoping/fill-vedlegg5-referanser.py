"""Vedlegg 5 — Referansebeskrivelser.

Malen har 5 identiske 6x2-tabeller (én per referanse). Vi fyller de tre vi har og
lar de to siste stå tomme (eller fjerner dem etter eget valg — la stå tomme i denne
versjonen så ADNO ser at vi prioriterer kvalitet, ikke kvantitet).
"""
from docx import Document
import sys
sys.path.insert(0, str(__import__("pathlib").Path(__file__).parent))
from _paths import SRC_DIR, DST_DIR
from _docx_utils import set_cell_text, fill_label_value

SRC = SRC_DIR / "1.5 Vedlegg 5 - Referansebeskrivelser.docx"
DST = DST_DIR / "Vedlegg 5 — Referansebeskrivelser (utfylt).docx"

REFERANSER = [
    {
        "tittel": "TryggDrift (Norsk Landbruksrådgiving) — HMS-app for landbruket",
        "oppdragsgiver": "Norsk Landbruksrådgiving (NLR)",
        "tidsrom": "2023–2024",
        "omfang": "Ca. 1 200 timer / ca. 1,5 MNOK ekskl. mva",
        "leveranse": """React Native-app for HMS-rapportering i landbruket. Asbjørn hadde teknisk lead-rolle. Erstattet eksisterende web-løsning med native app for både iOS og Android. Inkluderte regelverkdata, sjekklister, hendelsesregistrering og offline-modus. Levert og publisert i App Store og Google Play.

Nøkkelelementer:
- Regulert/samfunnskritisk domene (HMS-regelverk)
- Native publisering iOS + Android
- Offline-tilgang til regeldata
- Push-varsler ved nye hendelser/oppdateringer""",
        "relevans": """Nærmeste parallell til antidoping-appen vi har:
- Samme tekniske stakk (React Native)
- Samme leveranseform (App Store + Google Play, native)
- Samme type innhold (regelverk + sjekklister + ressurser for sluttbrukere)
- Samme type oppdragsgiver (uavhengig stiftelse / fagorgan)
- Asbjørn var teknisk lead — direkte overførbar erfaring""",
    },
    {
        "tittel": "Kolumbus mobil-app — kollektivtrafikk i Rogaland",
        "oppdragsgiver": "Kolumbus AS",
        "tidsrom": "2022–2025",
        "omfang": "Del av lengre engasjement, totalt ca. 2 000 timer dedikert utvikling",
        "leveranse": """Flutter-app for kollektivtrafikk med 142 000+ månedlige aktive brukere i Rogaland. Asbjørn bidro som senior utvikler i tverrfaglig team. Appen vant Nordic Public Transport Design Award 2025.

Funksjonalitet:
- Sanntids reiseinformasjon
- Push-varsler om endringer
- Universell utforming på høyt nivå
- Integrasjoner mot Entur og interne backend-tjenester
- Offline-modus for billett og reiseplan""",
        "relevans": """Viser at Asbjørn kan bygge en app med høy daglig bruk og høyt kvalitetsnivå:
- Stor brukerbase (142k+ MAU) — viser stabilitet og UX-modenhet
- Prisvinnende UU-implementasjon — direkte relevant for ADNO-appens UU-krav
- Push-varsler og kontekstuell informasjon — samme mekanikker som antidoping-appen
- Flutter-erfaring (alternativ til React Native ved behov)""",
    },
    {
        "tittel": "EaseePay (Easee) — betaling for EV-lading",
        "oppdragsgiver": "Easee AS",
        "tidsrom": "2022–2023",
        "omfang": "Ca. 1 000 timer / ca. 1,3 MNOK ekskl. mva",
        "leveranse": """Sikkerhetskritisk betalingsmodul i Easee-appen. Adyen-integrasjon, PCI DSS-samsvar, kryptering av betalingsdata, sertifikat-pinning, sikker token-håndtering. Asbjørn hadde teknisk hovedansvar for betalings­integrasjonen.

Inkluderte:
- PCI DSS-rammeverk for håndtering av kortdata
- Adyen Drop-in og 3DS-flyt
- Logging og overvåking via Sentry
- Sikker lagring av tokens (Keychain / Keystore)""",
        "relevans": """Demonstrerer evne til å håndtere regulatoriske krav og sensitive data — relevant for antidoping-appen:
- Personvern og dataminimering (samme prinsipper som GDPR-hensyn for utøverdata)
- Sertifikat-pinning og sikker kommunikasjon
- Kontroll over hva som logges og hvor
- Erfaring med strenge eksterne revisjonskrav (PCI DSS) — overførbar til helsekontekst""",
    },
]


def fill_referanse_table(table, ref: dict):
    set_cell_text(table.cell(0, 0), f"Beskrivelse av referanse: {ref['tittel']}", bold=True)
    set_cell_text(table.cell(0, 1), f"Beskrivelse av referanse: {ref['tittel']}", bold=True)
    fill_label_value(table, "Oppdragsgiver", ref["oppdragsgiver"])
    fill_label_value(table, "Tidsrom for gjennomføring", ref["tidsrom"])
    fill_label_value(table, "Omfang", ref["omfang"])
    fill_label_value(table, "Kort beskrivelse", ref["leveranse"])
    fill_label_value(table, "Relevans", ref["relevans"])


def main():
    doc = Document(SRC)

    # Malen har 5 referanse-tabeller. Fyll de tre vi har; la 4 og 5 stå tomme.
    for i, ref in enumerate(REFERANSER):
        if i < len(doc.tables):
            fill_referanse_table(doc.tables[i], ref)

    DST.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DST)
    print(f"✓ Lagret: {DST} ({len(REFERANSER)} referanser fylt)")
    print("  ⚠ Husk å innhente skriftlig samtykke fra kontaktpersoner før innsending.")


if __name__ == "__main__":
    main()

"""Hjelpere for docx-celle-utfylling — gjenbrukt på tvers av antidoping-scripts."""
import re
from docx.shared import Pt


def set_cell_text(cell, text: str, bold: bool = False, size: int = 10):
    """Tøm celle og sett ny tekst (multilinje støttet)."""
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""
    lines = text.split("\n")
    p0 = cell.paragraphs[0]
    run = p0.add_run(lines[0])
    run.bold = bold
    run.font.size = Pt(size)
    for line in lines[1:]:
        p = cell.add_paragraph()
        r = p.add_run(line)
        r.font.size = Pt(size)


# ---------- Markdown → docx renderer ----------

_INLINE_BOLD = re.compile(r"\*\*([^*]+)\*\*")
_INLINE_ITALIC = re.compile(r"(?<!\*)\*([^*\n]+)\*(?!\*)")
_INLINE_CODE = re.compile(r"`([^`]+)`")
_INLINE_LINK = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
_INLINE_STRIKE = re.compile(r"~~([^~]+)~~")


def _add_runs(paragraph, text: str, size: int = 10, base_bold: bool = False):
    """Skriv tekst med inline-markdown (bold/italic/code/link) som docx-runs."""
    # Reduser link til bare tekst
    text = _INLINE_LINK.sub(r"\1", text)
    # Strip strikethrough markup (vises som vanlig tekst)
    text = _INLINE_STRIKE.sub(r"\1", text)

    pos = 0
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|(?<!\*)\*[^*\n]+\*(?!\*))")
    for m in pattern.finditer(text):
        if m.start() > pos:
            r = paragraph.add_run(text[pos:m.start()])
            r.font.size = Pt(size)
            r.bold = base_bold
        token = m.group(0)
        r = paragraph.add_run()
        r.font.size = Pt(size)
        if token.startswith("**"):
            r.text = token[2:-2]
            r.bold = True
        elif token.startswith("`"):
            r.text = token[1:-1]
            r.font.name = "Consolas"
        else:  # *italic*
            r.text = token[1:-1]
            r.italic = True
            r.bold = base_bold
        pos = m.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        r.font.size = Pt(size)
        r.bold = base_bold


def _strip_md_inline(s: str) -> str:
    """Plain-tekst utgave (for tabellrendering)."""
    s = _INLINE_LINK.sub(r"\1", s)
    s = _INLINE_BOLD.sub(r"\1", s)
    s = _INLINE_ITALIC.sub(r"\1", s)
    s = _INLINE_CODE.sub(r"\1", s)
    s = _INLINE_STRIKE.sub(r"\1", s)
    return s


def render_markdown_to_cell(cell, md_text: str, size: int = 10):
    """Render markdown som formaterte paragraphs i en Word-celle.

    Støtter: #/##/### overskrifter, **bold**, *italic*, `code`, lister
    (- og 1.), tabeller (| a | b |), code blocks (``` ```), horisontal rule.
    """
    # Tøm celle
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""
    # Fjern alle eksisterende paragraphs unntatt første
    el = cell._tc
    for p in cell.paragraphs[1:]:
        el.remove(p._p)

    lines = md_text.split("\n")
    first = True

    def new_para():
        nonlocal first
        if first:
            first = False
            return cell.paragraphs[0]
        return cell.add_paragraph()

    i = 0
    in_code = False
    code_buf: list[str] = []
    table_buf: list[str] = []

    while i < len(lines):
        line = lines[i]

        # Code block
        if line.strip().startswith("```"):
            if in_code:
                # Lukk
                p = new_para()
                for j, cl in enumerate(code_buf):
                    if j > 0:
                        p = cell.add_paragraph()
                    r = p.add_run(cl)
                    r.font.size = Pt(size - 1)
                    r.font.name = "Consolas"
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # Tabell-buffer
        if line.lstrip().startswith("|") and "|" in line.lstrip()[1:]:
            table_buf.append(line)
            i += 1
            continue
        if table_buf:
            _render_md_table(cell, table_buf, new_para, size)
            table_buf = []
            # fall through til vanlig håndtering av denne linja

        # Horisontal rule
        if re.match(r"^---+$", line.strip()):
            p = new_para()
            r = p.add_run("─" * 30)
            r.font.size = Pt(size)
            i += 1
            continue

        # Overskrifter
        h = re.match(r"^(#{1,6})\s+(.*)$", line)
        if h:
            level = len(h.group(1))
            text = h.group(2).strip()
            p = new_para()
            r = p.add_run(text)
            r.bold = True
            r.font.size = Pt(size + max(0, 4 - level))
            i += 1
            continue

        # Bullet liste
        b = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if b:
            indent = len(b.group(1))
            p = new_para()
            r0 = p.add_run(("    " * (indent // 2)) + "• ")
            r0.font.size = Pt(size)
            _add_runs(p, b.group(2), size=size)
            i += 1
            continue

        # Nummerert liste
        n = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
        if n:
            indent = len(n.group(1))
            p = new_para()
            r0 = p.add_run(("    " * (indent // 2)) + f"{n.group(2)}. ")
            r0.font.size = Pt(size)
            _add_runs(p, n.group(3), size=size)
            i += 1
            continue

        # Tom linje
        if not line.strip():
            new_para()
            i += 1
            continue

        # Vanlig avsnitt
        p = new_para()
        _add_runs(p, line, size=size)
        i += 1

    if table_buf:
        _render_md_table(cell, table_buf, new_para, size)


def _render_md_table(cell, lines: list[str], new_para, size: int):
    """Render markdown-tabell som monospace-tekst i Word-cella.

    Nested tabeller i celler gir generelt dårlig layout — vi
    rendrer kompakt med | som separator i mono.
    """
    # Filtrer bort separator-rad (|---|---|)
    rows = []
    for ln in lines:
        s = ln.strip().strip("|").strip()
        if re.match(r"^[\s\-:|]+$", s):
            continue
        cells = [c.strip() for c in s.split("|")]
        rows.append([_strip_md_inline(c) for c in cells])
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    widths = [0] * n_cols
    for r in rows:
        for j, c in enumerate(r):
            widths[j] = max(widths[j], len(c))
    for idx, r in enumerate(rows):
        line = " | ".join((r[j] if j < len(r) else "").ljust(widths[j]) for j in range(n_cols))
        p = new_para()
        run = p.add_run(line)
        run.font.size = Pt(size - 1)
        run.font.name = "Consolas"
        if idx == 0:
            run.bold = True


def fill_label_value(table, label_match: str, value: str):
    """Fyll inn verdi i andre kolonne der første kolonne starter med label_match."""
    for row in table.rows:
        if len(row.cells) < 2:
            continue
        first = row.cells[0].text.strip()
        if first.lower().startswith(label_match.lower()) and not row.cells[1].text.strip():
            set_cell_text(row.cells[1], value)
            return True
    return False

"""Fyll ut Vedlegg 2 — Forpliktelseserklæring for UU-underleverandør.

Genereres bare hvis vi faktisk har en signerende underleverandør (Funka, MediaLT el.).
Verdiene under er PLACEHOLDER — bytt ut når avtale er på plass.
"""
from docx import Document
import sys
sys.path.insert(0, str(__import__("pathlib").Path(__file__).parent))
from _paths import SRC_DIR, DST_DIR, LEVERANDOR
from _docx_utils import fill_label_value

SRC = SRC_DIR / "1.2 Vedlegg 2 - Forpliktelseserklæring.docx"
DST = DST_DIR / "Vedlegg 2 — Forpliktelseserklæring (UU-underleverandør).docx"

# TODO: bytt ut med faktisk underleverandør når avtale er signert
SAMARBEIDSPARTNER = {
    "navn": "[TODO: Funka Nu AB / MediaLT / annen UU-leverandør]",
    "orgnr": "[TODO]",
    "postadresse": "[TODO]",
    "postnummer": "[TODO]",
    "sted": "[TODO]",
    "land": "[TODO]",
}


def main():
    doc = Document(SRC)

    # Tabell 0: Leverandøren
    t0 = doc.tables[0]
    fill_label_value(t0, "Navnet på virksomheten", LEVERANDOR["navn"])
    fill_label_value(t0, "Organisasjonsnummer", LEVERANDOR["orgnr"])
    fill_label_value(t0, "Postadresse", LEVERANDOR["postadresse"])
    fill_label_value(t0, "Postnummer", LEVERANDOR["postnummer"])
    fill_label_value(t0, "Sted", LEVERANDOR["sted"])
    fill_label_value(t0, "Land", LEVERANDOR["land"])

    # Tabell 1: Samarbeidspartner
    t1 = doc.tables[1]
    fill_label_value(t1, "Navnet på virksomheten", SAMARBEIDSPARTNER["navn"])
    fill_label_value(t1, "Organisasjonsnummer", SAMARBEIDSPARTNER["orgnr"])
    fill_label_value(t1, "Postadresse", SAMARBEIDSPARTNER["postadresse"])
    fill_label_value(t1, "Postnummer", SAMARBEIDSPARTNER["postnummer"])
    fill_label_value(t1, "Sted", SAMARBEIDSPARTNER["sted"])
    fill_label_value(t1, "Land", SAMARBEIDSPARTNER["land"])

    # Tabell 3 (Erklæring): sted/dato — signeres av samarbeidspartner manuelt
    t3 = doc.tables[3]
    fill_label_value(t3, "Sted og dato", "[fylles ut og signeres av samarbeidspartner]")

    DST.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DST)
    print(f"✓ Lagret: {DST}")
    print("  ⚠ Avhenger av at vi faktisk har en signerende UU-underleverandør.")
    print("    Hvis ikke: behold Vedlegg 2 ut av innsending og kryss 'En leverandør' i Vedlegg 1.")


if __name__ == "__main__":
    main()

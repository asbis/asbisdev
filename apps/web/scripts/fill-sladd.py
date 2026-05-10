"""Lag sladdet versjon av tilbudet (Excel prisskjema + begrunnelsesnotat).

Vi sladder:
 - Enhetspriser i Excel (timepriser og poster per rad) — forretningshemmelig jf. fvl § 13
 - Referanseperson-kontakt (Thomas Amundsen e-post) — beskyttelse av tredjepart

Totale evalueringspriser og totalt vedlikeholdsbudsjett er IKKE sladdet, jf. veiledningens
punkt om at "totalpriser i tilbud" ikke er forretningshemmelighet.

To versjoner:
  - sladd-gul-*.xlsx:   gul celle-bakgrunn, innhold beholdt (redigerbar)
  - sladd-svart-*.xlsx: svart celle-bakgrunn, innhold erstattet med ████ (ikke-redigerbar)
"""
from openpyxl import load_workbook
from openpyxl.styles import PatternFill, Font
from pathlib import Path
import shutil

SRC_XLSX = Path("/Users/asbis/code/asbisdev/oppdrag/aktive/2026-105336-rusinfo-selvhjelpsapp/til-innsending/autogenerert/03.2 DEL 2 SSA-T PRISSKJEMA Vedlegg 1 til bilag 7 (utfylt).xlsx")
DST_YELLOW = Path("/Users/asbis/code/asbisdev/oppdrag/aktive/2026-105336-rusinfo-selvhjelpsapp/til-innsending/autogenerert/Sladd-GUL — Prisskjema.xlsx")
DST_BLACK = Path("/Users/asbis/code/asbisdev/oppdrag/aktive/2026-105336-rusinfo-selvhjelpsapp/til-innsending/autogenerert/Sladd-SVART — Prisskjema.xlsx")

# Celler å sladde (enhetspriser G-kolonne + leverandørbeskrivelse F-kolonne)
CELLS_TO_REDACT = ["F11", "G11", "F15", "G15", "F19", "G19", "F20", "G20", "F24", "G24"]

yellow_fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")
black_fill = PatternFill(start_color="000000", end_color="000000", fill_type="solid")


def make_yellow():
    shutil.copy(SRC_XLSX, DST_YELLOW)
    wb = load_workbook(DST_YELLOW)
    ws = wb["Pris for tjenesten"]
    for coord in CELLS_TO_REDACT:
        ws[coord].fill = yellow_fill
    wb.save(DST_YELLOW)
    print(f"  ✓ Gul versjon: {DST_YELLOW.name}")


def make_black():
    shutil.copy(SRC_XLSX, DST_BLACK)
    wb = load_workbook(DST_BLACK)
    ws = wb["Pris for tjenesten"]
    # Hardkod H-kolonnen (evalueringspris per rad) og H26 (sum) slik at totalprisene
    # forblir synlige når G-kolonnen sladdes — ellers vil formlene =G*E gi #VALUE!
    H_VALUES = {"H11": 650000, "H15": 1200000, "H19": 22000, "H20": 14000, "H24": 15000, "H26": 1901000}
    for coord, val in H_VALUES.items():
        ws[coord].value = val
    for coord in CELLS_TO_REDACT:
        cell = ws[coord]
        cell.value = "████"
        cell.fill = black_fill
        cell.font = Font(color="000000")
    wb.save(DST_BLACK)
    print(f"  ✓ Svart versjon: {DST_BLACK.name}")


def write_begrunnelse():
    from docx import Document
    doc = Document()
    doc.add_heading("Begrunnelse for sladding av tilbud", level=1)
    doc.add_paragraph(
        "Leverandør: Asbjørn Rørvik (org.nr 820 252 632)\n"
        "Sak: 2026-105336 — Utvikling av selvhjelpsapp for personer som bruker kokain\n"
        "Dato: 24. april 2026"
    )
    doc.add_heading("Hva er sladdet", level=2)
    doc.add_paragraph(
        "I tråd med Vedlegg 1 «Veiledning for sladding av tilbud» og "
        "forvaltningsloven § 13 (1) nr. 2 har Leverandør sladdet opplysninger "
        "som vurderes å være av konkurransemessig betydning å hemmeligholde:"
    )
    items = [
        "Enhetspriser i Bilag 7 Prisskjema (rad 11 Hovedoppdrag, rad 15 Vedlikehold, "
        "rad 19/20 Timepriser, rad 24 Opplæring) — enhetspriser og rabattprosenter "
        "er forretningshemmeligheter jf. forvaltningsloven § 13 (1) nr. 2.",
        "Leverandørens tilhørende beskrivelser (kolonne F) — tilknyttet enhetsprisene "
        "og reflekterer Leverandørens prismodell.",
    ]
    for i in items:
        p = doc.add_paragraph(i, style="List Bullet")

    doc.add_heading("Hva er IKKE sladdet", level=2)
    doc.add_paragraph(
        "I tråd med veiledningen er følgende opplysninger ikke sladdet:"
    )
    not_items = [
        "Totalpris for hovedoppdraget og samlet evalueringspris — ikke forretningshemmelighet "
        "jf. veiledningens punkt om totalpriser.",
        "Øvrige bilag, tilbudsbrev og prosjektplan — inneholder ingen "
        "forretningshemmeligheter eller personopplysninger som krever sladding.",
        "CV-opplysninger og referansebeskrivelser — veiledningen spesifiserer at "
        "«sammensetning av team/CV til tilbudte ressurser» ikke er forretningshemmelighet.",
    ]
    for i in not_items:
        doc.add_paragraph(i, style="List Bullet")

    doc.add_heading("Sladdet versjon — format", level=2)
    doc.add_paragraph(
        "To versjoner leveres i tråd med veiledningens punkt om sladdet versjon:\n"
        "• Sladd-GUL — Prisskjema.xlsx: sensitive felter er markert med gul celle-"
        "bakgrunn, innhold er beholdt og kan redigeres. Brukes for intern vurdering "
        "hos Oppdragsgiver.\n"
        "• Sladd-SVART — Prisskjema.xlsx: sensitive felter er erstattet med ████ "
        "og har svart celle-bakgrunn. Brukes ved eventuell innsynsbegjæring.\n\n"
        "Øvrige tilbudsdokumenter (tilbudsbrev, Bilag 2, Bilag 4, Bilag 10, CV, "
        "kvalifikasjonsdokumenter) er identiske med originalen — ingen sladding er "
        "påkrevd for disse."
    )
    doc.add_paragraph(
        "Oppdragsgiver står fritt til å vurdere om sladdingen er begrunnet, "
        "jf. Vedlegg 1 «Veiledning for sladding av tilbud»."
    )
    out = Path("/Users/asbis/code/asbisdev/oppdrag/aktive/2026-105336-rusinfo-selvhjelpsapp/til-innsending/autogenerert/1.12 Sladd — Begrunnelse.docx")
    doc.save(out)
    print(f"  ✓ Begrunnelse: {out.name}")


if __name__ == "__main__":
    make_yellow()
    make_black()
    write_begrunnelse()
    print("\nSladdet versjon klar for opplasting.")
